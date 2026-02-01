import io
import cv2
import numpy as np
import fitz
from docx import Document
from openpyxl import load_workbook
from PIL import Image, ImageDraw, ImageFont, ImageFilter
import re
try:
    import pytesseract
except Exception:
    pytesseract = None
try:
    import phonenumbers
except Exception:
    phonenumbers = None


def redact_image_bytes(data: bytes, regions: list, mode: str = "blackout") -> bytes:
    arr = np.frombuffer(data, np.uint8)
    img = cv2.imdecode(arr, cv2.IMREAD_UNCHANGED)
    if img is None:
        raise ValueError("Unable to decode image")
    for r in regions:
        # r expected [x, y, w, h]
        x, y, w, h = r
        x, y, w, h = int(x), int(y), int(w), int(h)
        if mode == "blur":
            roi = img[y:y+h, x:x+w]
            k = max(3, (w//7)|1)
            roi = cv2.GaussianBlur(roi, (k, k), 0)
            img[y:y+h, x:x+w] = roi
        else:
            if img.ndim == 3:
                img[y:y+h, x:x+w] = (0, 0, 0)
            else:
                img[y:y+h, x:x+w] = 0
    ext = ".png"
    success, out = cv2.imencode(ext, img)
    if not success:
        raise ValueError("Failed to encode image")
    return out.tobytes()


def redact_pdf_bytes(data: bytes, regions: list) -> bytes:
    # regions: list of {"page": int, "rect": [x0,y0,x1,y1]}
    try:
        try:
            print(f"[redact_pdf_bytes] incoming regions type={type(regions)} count={len(regions) if regions else 0} preview_sample={regions[0] if regions else ''}")
        except Exception:
            print(f"[redact_pdf_bytes] incoming regions: {regions}")
        doc = fitz.open(stream=data, filetype="pdf")
        # Normalize regions if they are simple [x,y,w,h] canvas coords coming from the preview
        try:
            if isinstance(regions, list) and regions and isinstance(regions[0], (list, tuple)):
                zoom = 2.0
                page_pix_heights = []
                for pno in range(len(doc)):
                    pg = doc.load_page(pno)
                    pix = pg.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
                    page_pix_heights.append(pix.height)
                cum = [0]
                for hgt in page_pix_heights:
                    cum.append(cum[-1] + hgt)
                total_pages = len(page_pix_heights)
                norm = []
                for r in regions:
                    try:
                        x = float(r[0]); y = float(r[1]); w = float(r[2]); h = float(r[3])
                        pidx = 0
                        for i in range(total_pages):
                            if y >= cum[i] and y < cum[i+1]:
                                pidx = i; break
                        y_in_page = y - cum[pidx]
                        x0 = x/zoom; y0 = y_in_page/zoom; x1 = (x+w)/zoom; y1 = (y_in_page + h)/zoom
                        norm.append({"page": pidx, "rect": [x0, y0, x1, y1]})
                    except Exception:
                        continue
                regions = norm
        except Exception:
            pass
        # First, redact any email addresses by replacing them with a masked username and
        # redact phone numbers by drawing a black rectangle over their areas.
        try:
            for pno in range(len(doc)):
                page = doc.load_page(pno)
                text = page.get_text()
                # emails
                for m in EMAIL_RE.findall(text):
                    try:
                        if not isinstance(m, str):
                            try:
                                m = str(m)
                            except Exception:
                                print(f"[redact_pdf_bytes] coercion failed for email match type={type(m)} repr={m}")
                        areas = page.search_for(m)
                        for r in areas:
                            try:
                                # clear area then insert masked email (preserve domain)
                                shape = page.new_shape()
                                shape.draw_rect(r)
                                shape.finish(fill=(1, 1, 1))
                                shape.commit()
                                masked = mask_email_addr(m)
                                fontsize = max(6, (r.y1 - r.y0) * 0.7)
                                page.insert_textbox(r, masked, fontsize=fontsize, color=(0, 0, 0))
                            except Exception:
                                pass
                    except Exception as ex:
                        print(f"[redact_pdf_bytes] search_for failed for email match repr={repr(m)} type={type(m)} error={ex}")
                        continue
                # phone numbers -> black boxes
                for ph in PHONE_RE.findall(text):
                    try:
                        if not isinstance(ph, str):
                            try:
                                ph = str(ph)
                            except Exception:
                                print(f"[redact_pdf_bytes] coercion failed for phone match type={type(ph)} repr={ph}")
                        areas = page.search_for(ph)
                        for r in areas:
                            shape = page.new_shape()
                            shape.draw_rect(r)
                            shape.finish(fill=(0, 0, 0))
                            shape.commit()
                    except Exception as ex:
                        print(f"[redact_pdf_bytes] search_for failed for phone match repr={repr(ph)} type={type(ph)} error={ex}")
                        continue
        except Exception:
            pass
        for item in regions:
            # support both dict items and fallback list rects
            if isinstance(item, dict):
                page_no = int(item.get("page", 0))
                rect = item.get("rect")
            elif isinstance(item, (list, tuple)) and len(item) >= 4:
                page_no = 0
                rect = [item[0], item[1], item[0] + item[2], item[1] + item[3]]
            else:
                continue
            if rect is None:
                continue
            x0, y0, x1, y1 = map(float, rect)
            page = doc.load_page(page_no)
            shape = page.new_shape()
            shape.draw_rect(fitz.Rect(x0, y0, x1, y1))
            shape.finish(fill=(0, 0, 0))
            shape.commit()
        buf = io.BytesIO()
        doc.save(buf)
        out_bytes = buf.getvalue()
        try:
            modified_flag = (out_bytes != data)
            print(f"[redact_pdf_bytes] saved bytes len={len(out_bytes)} modified={modified_flag}")
        except Exception:
            print("[redact_pdf_bytes] saved bytes computed, but failed to compare to original")
        doc.close()
        return out_bytes
    except Exception as e:
        # Log the error and fall back to returning the original data so the user still
        # receives a downloadable PDF instead of causing a 500 with no download.
        try:
            import traceback
            print(f"[redact_pdf_bytes] ERROR: {e}\n" + traceback.format_exc())
        except Exception:
            print(f"[redact_pdf_bytes] ERROR: {e}")
        return data


def redact_docx_bytes(data: bytes, phrases: list, media_to_blur: list = None) -> bytes:
    buf = io.BytesIO(data)
    doc = Document(buf)
    def mask_text(s):
        return "█" * len(s)
    # replace email addresses and phone numbers across paragraphs
    try:
        for p in doc.paragraphs:
            text = p.text
            # mask emails (preserve domain)
            new_text = EMAIL_RE.sub(lambda m: mask_email_addr(m.group(0)), text)
            # replace phone numbers with black box characters
            new_text = PHONE_RE.sub(lambda m: '█' * len(m.group(0)), new_text)
            for ph in phrases:
                if ph in new_text:
                    new_text = new_text.replace(ph, mask_text(ph))
            if new_text != text:
                for r in list(p.runs):
                    r.text = ""
                p.add_run(new_text)
    except Exception:
        pass
    # also redact inside tables
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    new_text = EMAIL_RE.sub(lambda m: mask_email_addr(m.group(0)), text)
                    new_text = PHONE_RE.sub(lambda m: '█' * len(m.group(0)), new_text)
                    for ph in phrases:
                        if ph in new_text:
                            new_text = new_text.replace(ph, mask_text(ph))
                    if new_text != text:
                        for cp in list(cell.paragraphs):
                            for r in list(cp.runs):
                                r.text = ""
                        if cell.paragraphs:
                            cell.paragraphs[0].add_run(new_text)
    except Exception:
        pass
    out = io.BytesIO()
    doc.save(out)
    out_bytes = out.getvalue()
    # blur embedded images in the docx package (word/media)
    try:
        out_bytes = blur_media_in_ooxml(out_bytes, prefixes=('word/media/',), only_names=media_to_blur)
    except Exception:
        pass
    return out_bytes


def redact_xlsx_bytes(data: bytes, cells: list, columns: list, rows: list = None, phrases: list = None, media_to_blur: list = None) -> bytes:
    buf = io.BytesIO(data)
    wb = load_workbook(filename=buf)
    for ws in wb.worksheets:
        # first, redact email addresses in all string cells
        for row in ws.iter_rows(values_only=False):
            for cell in row:
                val = cell.value
                if isinstance(val, str) and EMAIL_RE.search(val):
                    try:
                        cell.value = EMAIL_RE.sub(lambda m: mask_email_addr(m.group(0)), val)
                    except Exception:
                        pass
                if isinstance(val, str) and PHONE_RE.search(val):
                    try:
                        # replace phone occurrences with black box characters
                        cell.value = PHONE_RE.sub(lambda m: '█' * len(m.group(0)), cell.value)
                    except Exception:
                        pass
                # redact any user-specified phrases inside cells
                if phrases and isinstance(val, str):
                    try:
                        for ph in phrases:
                            if not ph: continue
                            if ph in cell.value:
                                repl = '█' * len(ph)
                                cell.value = cell.value.replace(ph, repl)
                    except Exception:
                        pass
        # mask specific cells like "A1", "B2"
        for c in cells:
            try:
                    cell = ws[c]
                    val = cell.value
                    if isinstance(val, str):
                        # replace text with black box characters of same length
                        cell.value = '█' * len(val)
                    else:
                        cell.value = "REDACTED"
            except Exception:
                continue
        # mask entire columns like "C" or index
        for col in columns:
                for row_cells in ws.iter_rows(min_row=1, max_row=ws.max_row):
                    try:
                        if isinstance(col, int):
                            cell = row_cells[col-1]
                        else:
                            # assume column letter; map to the current row
                            cell = ws[f"{col}{row_cells[0].row}"]
                        val = cell.value
                        if isinstance(val, str):
                            cell.value = '█' * len(val)
                        else:
                            cell.value = "REDACTED"
                    except Exception:
                        continue
        # mask entire rows (rows may be list of ints or ranges)
        if rows:
            for r in rows:
                try:
                    if isinstance(r, list) or isinstance(r, tuple):
                        # treat as [start,end]
                        start = int(r[0]); end = int(r[1])
                        rng = range(max(1, start), min(ws.max_row, end) + 1)
                    else:
                        rng = [int(r)]
                    for rr in rng:
                        for c in range(1, ws.max_column + 1):
                            cell = ws.cell(row=rr, column=c)
                            val = cell.value
                            if isinstance(val, str):
                                cell.value = '█' * len(val)
                            else:
                                cell.value = "REDACTED"
                except Exception:
                    continue
    out = io.BytesIO()
    wb.save(out)
    out_bytes = out.getvalue()
    # blur embedded images in the xlsx package (xl/media)
    try:
        out_bytes = blur_media_in_ooxml(out_bytes, prefixes=('xl/media/',), only_names=media_to_blur)
    except Exception:
        pass
    return out_bytes


EMAIL_RE = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")
PHONE_RE = re.compile(r"\+?\d[\d\s\-\(\)]{6,}\d")
# E.164-like generic phone pattern (validate with context unless leading +)
E164_RE = re.compile(r"\+?[1-9]\d{6,14}\b")
# account numeric-only (8-16 digits) and generic allowing separators
ACC_RE = re.compile(r"\b\d{8,16}\b")
ACC_GENERIC_RE = re.compile(r"\b(?:\d[ \-]?){8,16}\b")
# sequences of 7+ digits (catch common phone formats without punctuation)
DIGITSEQ_RE = re.compile(r"\b\d{7,}\b")

# Additional detection rules
CARD_RE = re.compile(r"\b(?:\d[ -]*?){13,19}\b")
IBAN_RE = re.compile(r"\b[A-Z]{2}\d{2}[A-Z0-9]{4,30}\b")
SWIFT_RE = re.compile(r"\b[A-Z]{6}[A-Z0-9]{2}(?:[A-Z0-9]{3})?\b")
API_KEY_RE = re.compile(r"\b(?:AKIA[0-9A-Z]{16}|[A-Za-z0-9_\-]{32,64})\b")

# Routing / bank-specific patterns
ROUTING_RE = re.compile(r"\b\d{9}\b")
SORTCODE_RE = re.compile(r"\b\d{2}[- ]?\d{2}[- ]?\d{2}\b")

# More sensitive patterns
SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
SSN_RE2 = re.compile(r"\b\d{9}\b")
PAN_RE = re.compile(r"\b[A-Z]{5}\d{4}[A-Z]\b")
AADHAAR_RE = re.compile(r"\b\d{4}[\s-]?\d{4}[\s-]?\d{4}\b")
IFSC_RE = re.compile(r"\b[A-Z]{4}0[A-Z0-9]{6}\b", re.IGNORECASE)
MASKED_CARD_RE = re.compile(r"(?:\*{4}[ \-]?){1,3}\d{4}")
CVV_RE = re.compile(r"\b\d{3,4}\b")
PIN_RE = re.compile(r"\b\d{4,6}\b")
JWT_RE = re.compile(r"\beyJ[A-Za-z0-9_\-]+\.[A-Za-z0-9_\-]+\.[A-Za-z0-9_\-]+\b")

# OTP / 2FA codes (detect with context)
OTP_RE = re.compile(r"\b\d{4,8}\b")

# Biometric/passport/license heuristics (keywords + short id patterns)
BIOMETRIC_KEYWORDS = ['fingerprint', 'retina', 'iris', 'voice', 'dna', 'facial', 'face recognition', 'biometric']
PASSPORT_RE = re.compile(r"\b[A-Z0-9]{6,9}\b")
DL_RE = re.compile(r"\b[A-Z0-9\-]{5,12}\b")


def _shannon_entropy(s: str) -> float:
    # simple entropy estimator for strings
    from math import log2
    if not s:
        return 0.0
    probs = [float(s.count(c)) / len(s) for c in set(s)]
    return -sum(p * log2(p) for p in probs)


def _digits_only(s: str) -> str:
    return re.sub(r"\D", "", s)


def luhn_check(num: str) -> bool:
    s = _digits_only(num)
    if len(s) < 13 or len(s) > 19:
        return False
    total = 0
    alt = False
    for ch in s[::-1]:
        d = ord(ch) - 48
        if alt:
            d = d * 2
            if d > 9:
                d -= 9
        total += d
        alt = not alt
    return (total % 10) == 0


def iban_check(iban: str) -> bool:
    try:
        s = iban.replace(' ', '').upper()
        if len(s) < 5:
            return False
        # move first four chars to end
        rearr = s[4:] + s[:4]
        # convert letters to numbers A=10 ... Z=35
        conv = ''
        for ch in rearr:
            if ch.isalpha():
                conv += str(ord(ch) - 55)
            else:
                conv += ch
        # perform mod 97 (do in chunks)
        remainder = 0
        for i in range(0, len(conv), 9):
            part = str(remainder) + conv[i:i+9]
            remainder = int(part) % 97
        return remainder == 1
    except Exception:
        return False


def verhoeff_check(num: str) -> bool:
    """Validate a numeric string using the Verhoeff algorithm (used for Aadhaar)."""
    try:
        s = _digits_only(num)
        if not s or not s.isdigit():
            return False
        # multiplication table d
        d_table = [
            [0,1,2,3,4,5,6,7,8,9],
            [1,2,3,4,0,6,7,8,9,5],
            [2,3,4,0,1,7,8,9,5,6],
            [3,4,0,1,2,8,9,5,6,7],
            [4,0,1,2,3,9,5,6,7,8],
            [5,9,8,7,6,0,4,3,2,1],
            [6,5,9,8,7,1,0,4,3,2],
            [7,6,5,9,8,2,1,0,4,3],
            [8,7,6,5,9,3,2,1,0,4],
            [9,8,7,6,5,4,3,2,1,0]
        ]
        # permutation table p
        p_table = [
            [0,1,2,3,4,5,6,7,8,9],
            [1,5,7,6,2,8,3,0,9,4],
            [5,8,0,3,7,9,6,1,4,2],
            [8,9,1,6,0,4,3,5,2,7],
            [9,4,5,3,1,2,6,8,7,0],
            [4,2,8,6,5,7,3,9,0,1],
            [2,7,9,3,8,0,6,4,1,5],
            [7,0,4,6,9,1,3,2,5,8]
        ]
        # inverse table
        inv = [0,4,3,2,1,5,6,7,8,9]
        c = 0
        # process digits from right to left
        for i, ch in enumerate(reversed(s)):
            c = d_table[c][p_table[i % 8][int(ch)]]
        return c == 0
    except Exception:
        return False


_CONTEXT_TOKENS = {
    'account': ['account', 'acct', 'iban', 'routing', 'bank', 'acct no', 'account no', 'accno'],
    'card': ['card', 'credit', 'visa', 'mastercard', 'amex', 'cvv', 'card number'],
    'ssn': ['ssn', 'social security', 'social'],
    'credentials': ['password', 'pwd', 'secret', 'token', 'api_key', 'apikey']
}

# extend context tokens for OTP, routing, and credentials
_CONTEXT_TOKENS['otp'] = ['otp', 'one-time', '2fa', 'two-factor', 'verification', 'code']
_CONTEXT_TOKENS['routing'] = ['routing', 'aba', 'routing number', 'sort code']
_CONTEXT_TOKENS['biometric'] = ['fingerprint', 'retina', 'iris', 'voice', 'dna', 'biometric', 'facial']
_CONTEXT_TOKENS['phone'] = ['phone', 'tel', 'mobile', 'msisdn', 'contact']
_CONTEXT_TOKENS['aadhaar'] = ['aadhaar', 'aadhar', 'uid']


def _has_context(text: str, start: int, end: int, tokens: list, window: int = 80) -> bool:
    lo = max(0, start - window)
    hi = min(len(text), end + window)
    ctx = text[lo:hi].lower()
    for t in tokens:
        if t in ctx:
            return True
    return False


def scan_text_for_sensitive_data(text: str, require_context_for=('acc',)) -> list:
    """Return list of matches: {'category','match','start','end'}"""
    out = []
    if not text:
        return out
    # emails
    for m in EMAIL_RE.finditer(text):
        out.append({'category': 'email', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # phones
    for m in PHONE_RE.finditer(text):
        # include obvious phone matches (contain separators or leading +) or require context
        s = m.group(0)
        if s.strip().startswith('+') or _has_context(text, m.start(), m.end(), _CONTEXT_TOKENS.get('phone', []), window=40):
            out.append({'category': 'phone', 'match': s, 'start': m.start(), 'end': m.end()})
    # also catch E.164-like sequences but require context unless a + is present
    for m in E164_RE.finditer(text):
        s = m.group(0)
        if s.startswith('+') or _has_context(text, m.start(), m.end(), _CONTEXT_TOKENS.get('phone', []), window=40):
            out.append({'category': 'phone', 'match': s, 'start': m.start(), 'end': m.end()})
    # IBAN
    for m in IBAN_RE.finditer(text):
        if iban_check(m.group(0)):
            out.append({'category': 'iban', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # credit cards (use Luhn)
    for m in CARD_RE.finditer(text):
        candidate = _digits_only(m.group(0))
        if luhn_check(candidate):
            out.append({'category': 'credit_card', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # SSN (US) patterns
    for m in SSN_RE.finditer(text):
        out.append({'category': 'ssn', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    for m in SSN_RE2.finditer(text):
        # numeric-only 9-digit could be SSN but check context
        if _has_context(text, m.start(), m.end(), ['ssn', 'social security', 'ssn:'], window=50):
            out.append({'category': 'ssn', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # PAN (India)
    for m in PAN_RE.finditer(text):
        out.append({'category': 'pan', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # Aadhaar (India) 12-digit - validate with Verhoeff checksum when possible
    for m in AADHAAR_RE.finditer(text):
        digits = _digits_only(m.group(0))
        if verhoeff_check(digits):
            out.append({'category': 'aadhaar', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # CVV / PIN: only include when context nearby mentions cvv/pin/cvc
    for m in CVV_RE.finditer(text):
        if _has_context(text, m.start(), m.end(), ['cvv', 'cvc', 'security code', 'cvv:'], window=20):
            out.append({'category': 'cvv', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    for m in PIN_RE.finditer(text):
        if _has_context(text, m.start(), m.end(), ['pin', 'passcode', 'pin:'], window=20):
            out.append({'category': 'pin', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # masked card numbers like **** **** **** 1234
    for m in MASKED_CARD_RE.finditer(text):
        out.append({'category': 'credit_card_masked', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # JWT-like tokens and high-entropy strings (credentials/secrets)
    for m in JWT_RE.finditer(text):
        out.append({'category': 'jwt', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    for m in API_KEY_RE.finditer(text):
        tok = m.group(0)
        if len(tok) >= 20 or _shannon_entropy(tok) > 4.0:
            out.append({'category': 'api_key', 'match': tok, 'start': m.start(), 'end': m.end()})
    # Credentials: look for password patterns like 'password: xyz' or 'pwd='
    for m in re.finditer(r"(?:password|pwd|pass|secret)\s*[:=]\s*(\S{4,})", text, re.IGNORECASE):
        out.append({'category': 'credential', 'match': m.group(1), 'start': m.start(1), 'end': m.end(1)})
    # Health / PHI heuristics: keyword detection
    health_keywords = ['medical', 'patient', 'diagnosis', 'prescription', 'mrn', 'medical record', 'lab result']
    for hk in health_keywords:
        for m in re.finditer(re.escape(hk), text, re.IGNORECASE):
            # include surrounding line
            lo = max(0, text.rfind('\n', 0, m.start()))
            hi = text.find('\n', m.end())
            if hi == -1: hi = min(len(text), m.end() + 120)
            snippet = text[lo:hi].strip()
            out.append({'category': 'health_info', 'match': snippet, 'start': lo, 'end': hi})
    # account-like numbers (require nearby context to reduce false positives)
    for pat in (ACC_RE, ACC_GENERIC_RE):
        for m in pat.finditer(text):
            if _has_context(text, m.start(), m.end(), _CONTEXT_TOKENS.get('account', []), window=60):
                out.append({'category': 'account', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # routing numbers (9-digit) - require context
    for m in ROUTING_RE.finditer(text):
        if _has_context(text, m.start(), m.end(), _CONTEXT_TOKENS.get('routing', []), window=60):
            out.append({'category': 'routing', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # sort codes (UK) - require context
    for m in SORTCODE_RE.finditer(text):
        if _has_context(text, m.start(), m.end(), _CONTEXT_TOKENS.get('routing', []), window=60):
            out.append({'category': 'sort_code', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # generic digit sequences (long) - include only with context
    for m in DIGITSEQ_RE.finditer(text):
        if _has_context(text, m.start(), m.end(), _CONTEXT_TOKENS.get('account', []), window=60):
            out.append({'category': 'digit_sequence', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # API keys / high-entropy tokens (simple heuristic)
    for m in API_KEY_RE.finditer(text):
        tok = m.group(0)
        if len(tok) >= 20:
            out.append({'category': 'api_key', 'match': tok, 'start': m.start(), 'end': m.end()})
    # OTP / two-factor codes: only include when nearby context suggests OTP/2FA
    for m in OTP_RE.finditer(text):
        if _has_context(text, m.start(), m.end(), _CONTEXT_TOKENS.get('otp', []), window=30):
            out.append({'category': 'otp', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # SWIFT/BIC
    for m in SWIFT_RE.finditer(text):
        out.append({'category': 'swift', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # IFSC / bank codes (India)
    for m in IFSC_RE.finditer(text):
        out.append({'category': 'ifsc', 'match': m.group(0), 'start': m.start(), 'end': m.end()})
    # biometric / passport / driver's license heuristics
    for kw in BIOMETRIC_KEYWORDS:
        for m in re.finditer(re.escape(kw), text, re.IGNORECASE):
            lo = max(0, text.rfind('\n', 0, m.start()))
            hi = text.find('\n', m.end())
            if hi == -1: hi = min(len(text), m.end() + 120)
            snippet = text[lo:hi].strip()
            out.append({'category': 'biometric', 'match': snippet, 'start': lo, 'end': hi})
    # passports / driver's license: look for nearby keywords and an alnum token
    for m in re.finditer(r"passport\s*[:#\-]?\s*([A-Z0-9\-]{6,9})", text, re.IGNORECASE):
        out.append({'category': 'passport', 'match': m.group(1), 'start': m.start(1), 'end': m.end(1)})
    for m in re.finditer(r"driver(?:'s)?\s+licen[cs]e\s*[:#\-]?\s*([A-Z0-9\-]{5,12})", text, re.IGNORECASE):
        out.append({'category': 'driver_license', 'match': m.group(1), 'start': m.start(1), 'end': m.end(1)})
    return out


def mask_email_addr(s: str) -> str:
    try:
        parts = s.split('@', 1)
        if len(parts) != 2:
            return s
        local, domain = parts
        if not local:
            return s
        return ('*' * max(3, len(local))) + '@' + domain
    except Exception:
        return s


def preview_pdf_first_page(data: bytes, zoom: float = 2.0) -> bytes:
    # Render all pages and concatenate vertically into one PNG
    doc = fitz.open(stream=data, filetype="pdf")
    images = []
    try:
        for pno in range(len(doc)):
            page = doc.load_page(pno)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
            images.append(img)
        if not images:
            return b""
        # concatenate vertically
        widths = [im.width for im in images]
        heights = [im.height for im in images]
        maxw = max(widths)
        totalh = sum(heights)
        out_img = Image.new('RGB', (maxw, totalh), color='white')
        y = 0
        for im in images:
            out_img.paste(im, (0, y))
            y += im.height
        out = io.BytesIO()
        out_img.save(out, format='PNG')
        return out.getvalue()
    finally:
        doc.close()


def detect_pdf_bytes(data: bytes):
    doc = fitz.open(stream=data, filetype="pdf")
    results = []
    for pno in range(len(doc)):
        page = doc.load_page(pno)
        text = page.get_text()
        matches = []
        for m in scan_text_for_sensitive_data(text):
            txt = m.get('match')
            try:
                areas = page.search_for(txt)
                if areas:
                    for r in areas:
                        matches.append({"text": txt, "rect": [r.x0, r.y0, r.x1, r.y1], "category": m.get('category')})
                else:
                    # fallback: use word boxes and sliding-window to find multi-word matches
                    try:
                        words = page.get_text("words")
                        import re as _re
                        phrase_norm = _re.sub(r"\s+", " ", _re.sub(r"[^\w\s]", " ", txt)).strip().lower()
                        if words and phrase_norm:
                            norm_words = [_re.sub(r"[^\w]", "", w[4]).lower() for w in words]
                            pw = [pw for pw in phrase_norm.split() if pw]
                            plen = len(pw)
                            if plen > 0:
                                for i in range(0, max(0, len(norm_words) - 0)):
                                    found_any = False
                                    for end in range(i + 1, min(len(norm_words), i + plen + 6) + 1):
                                        seq = norm_words[i:end]
                                        joined = " ".join(seq)
                                        if pw == seq or phrase_norm in joined or joined in phrase_norm:
                                            try:
                                                x0 = min(words[k][0] for k in range(i, end))
                                                y0 = min(words[k][1] for k in range(i, end))
                                                x1 = max(words[k][2] for k in range(i, end))
                                                y1 = max(words[k][3] for k in range(i, end))
                                                matches.append({"text": txt, "rect": [x0, y0, x1, y1], "category": m.get('category')})
                                                found_any = True
                                                break
                                            except Exception:
                                                continue
                                    if found_any:
                                        break
                    except Exception:
                        pass
            except Exception:
                continue
        if matches:
            results.append({"page": pno, "matches": matches})
    doc.close()
    return results


def detect_docx_bytes(data: bytes):
    buf = io.BytesIO(data)
    doc = Document(buf)
    found = []
    # check paragraphs
    for p in doc.paragraphs:
        t = p.text
        for m in scan_text_for_sensitive_data(t):
            found.append({'match': m.get('match'), 'category': m.get('category')})
    # check tables (cells)
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text
                    for m in scan_text_for_sensitive_data(t):
                        found.append({'match': m.get('match'), 'category': m.get('category')})
    except Exception:
        pass
    # also report embedded images (filenames) if present
    imgs = []
    img_matches = []
    try:
        # inspect package for word/media
        import zipfile
        z = zipfile.ZipFile(io.BytesIO(data))
        for name in z.namelist():
            if name.startswith('word/media/'):
                base = name.split('/')[-1]
                imgs.append(base)
                # attempt OCR on embedded image if pytesseract available
                try:
                    raw = z.read(name)
                    if pytesseract is not None:
                        det = detect_image_bytes(raw)
                        if isinstance(det, dict):
                            matches = det.get('matches', [])
                            ft = det.get('full_text', '')
                        else:
                            matches = det
                            ft = ''
                        if matches:
                            img_matches.append({'image': base, 'matches': matches, 'full_text': ft})
                except Exception:
                    pass
        z.close()
    except Exception:
        pass
    # dedupe while preserving order
    seen = set()
    unique = []
    for it in found:
        try:
            if isinstance(it, dict):
                key = (it.get('match'), it.get('category'))
            else:
                key = (it, None)
        except Exception:
            key = (str(it), None)
        if key not in seen:
            seen.add(key)
            unique.append(it)
    return {'text_matches': unique, 'images': imgs, 'image_matches': img_matches}


def detect_xlsx_bytes(data: bytes):
    buf = io.BytesIO(data)
    wb = load_workbook(filename=buf)
    found = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=False):
            for cell in row:
                val = cell.value
                if not isinstance(val, str):
                    continue
                for m in scan_text_for_sensitive_data(val):
                    found.append({"sheet": ws.title, "cell": cell.coordinate, "match": m.get('match'), 'category': m.get('category')})
    # also list images in xl/media
    imgs = []
    img_matches = []
    try:
        import zipfile
        z = zipfile.ZipFile(io.BytesIO(data))
        for name in z.namelist():
            if name.startswith('xl/media/'):
                base = name.split('/')[-1]
                imgs.append(base)
                try:
                    raw = z.read(name)
                    if pytesseract is not None:
                        det = detect_image_bytes(raw)
                        if isinstance(det, dict):
                            matches = det.get('matches', [])
                            ft = det.get('full_text', '')
                        else:
                            matches = det
                            ft = ''
                        if matches:
                            img_matches.append({'image': base, 'matches': matches, 'full_text': ft})
                except Exception:
                    pass
        z.close()
    except Exception:
        pass
    return {'text_matches': found, 'images': imgs, 'image_matches': img_matches}


def blur_media_in_ooxml(data: bytes, prefixes=('word/media/', 'xl/media/'), only_names: list = None) -> bytes:
    """Open OOXML package bytes, blur images under given prefixes, and return new package bytes.
    If `only_names` is provided, only those media file basenames will be blurred; others are preserved.
    """
    import zipfile
    inbuf = io.BytesIO(data)
    outbuf = io.BytesIO()
    with zipfile.ZipFile(inbuf, 'r') as zin:
        with zipfile.ZipFile(outbuf, 'w', compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                name = item.filename
                raw = zin.read(name)
                # blur if path matches prefixes and file looks like an image
                if any(name.startswith(p) for p in prefixes):
                    # if only_names provided, skip files not in list
                    if only_names is not None:
                        base = name.split('/')[-1]
                        if base not in only_names:
                            zout.writestr(name, raw)
                            continue
                    try:
                        im = Image.open(io.BytesIO(raw)).convert('RGBA')
                        im = im.filter(ImageFilter.GaussianBlur(radius=8))
                        wb = io.BytesIO()
                        # preserve original format
                        fmt = im.format or ('PNG' if name.lower().endswith('.png') else 'JPEG')
                        im.save(wb, format=fmt)
                        raw = wb.getvalue()
                    except Exception:
                        pass
                zout.writestr(name, raw)
    return outbuf.getvalue()


def detect_image_bytes(data: bytes):
    if pytesseract is None:
        return {"error": "pytesseract not installed"}
    arr = np.frombuffer(data, np.uint8)
    img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
    if img is None:
        return {"error": "cannot decode image"}
    # use pytesseract to get word boxes and full text
    d = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
    matches = []
    words = []
    n = len(d.get('text', []))
    for i in range(n):
        txt = (d['text'][i] or '').strip()
        if not txt:
            continue
        x = int(d.get('left', [0])[i]); y = int(d.get('top', [0])[i]); w = int(d.get('width', [0])[i]); h = int(d.get('height', [0])[i])
        words.append({'text': txt, 'x': x, 'y': y, 'w': w, 'h': h})
    # build full text (words separated by spaces)
    full_text = ' '.join([w['text'] for w in words])
    # scan for sensitive items in the full text
    found = scan_text_for_sensitive_data(full_text)
    # map each found match back to word boxes by searching words sequence
    for f in found:
        mtxt = f.get('match')
        if not mtxt:
            continue
        mlow = mtxt.lower()
        # try to find contiguous sequence of words that contain the match
        joined = ' '.join([w['text'] for w in words]).lower()
        idx = joined.find(mlow)
        if idx == -1:
            # fallback: try to find single word containing substring
            for w in words:
                if mlow in w['text'].lower():
                    matches.append({'text': mtxt, 'rect': [w['x'], w['y'], w['w'], w['h']], 'category': f.get('category')})
                    break
            continue
        # map character index back to word indices
        # compute cumulative lengths
        cum = []
        s = ''
        for wi, w in enumerate(words):
            if s:
                s += ' ' + w['text']
            else:
                s = w['text']
            cum.append(len(s))
        # find start/end word indices covering idx..idx+len(mlow)-1
        start_char = idx; end_char = idx + len(mlow) - 1
        start_w = 0
        while start_w < len(cum) and cum[start_w] <= start_char:
            start_w += 1
        end_w = start_w
        while end_w < len(cum) and cum[end_w] <= end_char:
            end_w += 1
        # build bounding box across words start_w..end_w
        try:
            xs = [words[k]['x'] for k in range(start_w, end_w+1)]
            ys = [words[k]['y'] for k in range(start_w, end_w+1)]
            xe = [words[k]['x'] + words[k]['w'] for k in range(start_w, end_w+1)]
            ye = [words[k]['y'] + words[k]['h'] for k in range(start_w, end_w+1)]
            x0 = min(xs); y0 = min(ys); x1 = max(xe); y1 = max(ye)
            matches.append({'text': mtxt, 'rect': [int(x0), int(y0), int(x1 - x0), int(y1 - y0)], 'category': f.get('category')})
        except Exception:
            # last-resort: single word fallback
            for w in words:
                if mlow in w['text'].lower():
                    matches.append({'text': mtxt, 'rect': [w['x'], w['y'], w['w'], w['h']], 'category': f.get('category')})
                    break
    return {'matches': matches, 'full_text': full_text}


def preview_docx_bytes(data: bytes, width: int = 800, line_height: int = 18) -> bytes:
    buf = io.BytesIO(data)
    doc = Document(buf)
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            # mask emails and phones for preview
            text = EMAIL_RE.sub(lambda m: mask_email_addr(m.group(0)), text)
            text = PHONE_RE.sub(lambda m: '█' * len(m.group(0)), text)
            lines.append(text)
    # include table cell contents for preview
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text = EMAIL_RE.sub(lambda m: mask_email_addr(m.group(0)), text)
                        text = PHONE_RE.sub(lambda m: '█' * len(m.group(0)), text)
                        lines.append(text)
    except Exception:
        pass
    # create image and draw text
    font = None
    try:
        font = ImageFont.load_default()
    except Exception:
        font = None
    img_h = max(200, line_height * (len(lines) + 2))
    out_img = Image.new('RGB', (width, img_h), color='white')
    draw = ImageDraw.Draw(out_img)
    y = 10
    for l in lines:
        draw.text((10, y), l, fill='black', font=font)
        y += line_height
    out = io.BytesIO()
    out_img.save(out, format='PNG')
    return out.getvalue()


def preview_docx_html(data: bytes) -> str:
    """Return a simple HTML representation of the DOCX with paragraphs and tables."""
    buf = io.BytesIO(data)
    doc = Document(buf)
    parts = []
    parts.append('<div style="font-family:Arial,Helvetica,sans-serif;color:#111">')
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        t = EMAIL_RE.sub(lambda m: mask_email_addr(m.group(0)), t)
        t = PHONE_RE.sub(lambda m: '█' * len(m.group(0)), t)
        parts.append(f'<p style="margin:6px 0;">{_escape_html(t)}</p>')
    # tables
    try:
        for table in doc.tables:
            parts.append('<table border="1" cellpadding="4" style="border-collapse:collapse;margin:8px 0;">')
            for row in table.rows:
                parts.append('<tr>')
                for cell in row.cells:
                    txt = cell.text.strip()
                    txt = EMAIL_RE.sub(lambda m: mask_email_addr(m.group(0)), txt)
                    txt = PHONE_RE.sub(lambda m: '█' * len(m.group(0)), txt)
                    parts.append(f'<td>{_escape_html(txt)}</td>')
                parts.append('</tr>')
            parts.append('</table>')
    except Exception:
        pass
    parts.append('</div>')
    return '\n'.join(parts)


def preview_xlsx_html(data: bytes, max_rows: int = None, max_cols: int = None) -> str:
    """Return an HTML table representation of the first sheet of the workbook.
    By default it will include all rows and columns unless `max_rows`/`max_cols` are provided.
    """
    buf = io.BytesIO(data)
    wb = load_workbook(filename=buf, data_only=True)
    ws = wb.active
    rows = ws.max_row if (max_rows is None) else min(ws.max_row, max_rows)
    cols = ws.max_column if (max_cols is None) else min(ws.max_column, max_cols)
    parts = []
    parts.append('<div style="overflow:auto;max-height:600px;border:1px solid #ddd;padding:6px;background:#fff">')
    parts.append('<table border="1" cellpadding="6" style="border-collapse:collapse;font-family:Arial,Helvetica,sans-serif;font-size:13px">')
    # header row with column letters
    parts.append('<thead><tr><th style="background:#f3f3f3"></th>')
    import string
    def col_letter(n):
        # 1-indexed
        s = ''
        while n>0:
            n, rem = divmod(n-1, 26)
            s = chr(65+rem) + s
        return s
    for c in range(1, cols+1):
        parts.append(f'<th style="background:#f6f8fa">{col_letter(c)}</th>')
    parts.append('</tr></thead>')
    parts.append('<tbody>')
    for r in range(1, rows+1):
        parts.append(f'<tr><th style="background:#f6f8fa">{r}</th>')
        for c in range(1, cols+1):
            val = ws.cell(row=r, column=c).value
            s = '' if val is None else str(val)
            # mask emails and phone-like sequences for preview
            s = EMAIL_RE.sub(lambda m: mask_email_addr(m.group(0)), s)
            s = PHONE_RE.sub(lambda m: '█' * len(m.group(0)), s)
            s = ACC_GENERIC_RE.sub(lambda m: '█' * len(m.group(0)), s)
            s = ACC_RE.sub(lambda m: '█' * len(m.group(0)), s)
            s = DIGITSEQ_RE.sub(lambda m: '█' * len(m.group(0)), s)
            parts.append(f'<td>{_escape_html(s)}</td>')
        parts.append('</tr>')
    parts.append('</tbody></table></div>')
    return '\n'.join(parts)


def _escape_html(s: str) -> str:
    import html
    return html.escape(s)


def preview_xlsx_bytes(data: bytes, width: int = 800, row_h: int = 24) -> bytes:
    buf = io.BytesIO(data)
    wb = load_workbook(filename=buf, data_only=True)
    ws = wb.active
    # collect cell values for first 10 rows/10 cols
    max_rows = min(20, ws.max_row)
    max_cols = min(10, ws.max_column)
    cell_texts = []
    # redact emails and phone numbers for preview
    for row in ws.iter_rows(min_row=1, max_row=max_rows, min_col=1, max_col=max_cols):
        for cell in row:
            val = cell.value
            if isinstance(val, str) and EMAIL_RE.search(val):
                cell.value = EMAIL_RE.sub(lambda m: mask_email_addr(m.group(0)), val)
            if isinstance(val, str) and PHONE_RE.search(val):
                cell.value = PHONE_RE.sub(lambda m: '█' * len(m.group(0)), cell.value)
    col_width = max(60, width // max_cols)
    img_h = max(200, row_h * (max_rows + 1))
    out_img = Image.new('RGB', (col_width * max_cols, img_h), color='white')
    draw = ImageDraw.Draw(out_img)
    font = None
    try:
        font = ImageFont.load_default()
    except Exception:
        font = None
    # header
    y = 0
    for r in range(1, max_rows + 1):
        x = 0
        for c in range(1, max_cols + 1):
            val = ws.cell(row=r, column=c).value
            s = '' if val is None else str(val)
            draw.rectangle([x, y, x + col_width - 1, y + row_h - 1], outline='gray')
            draw.text((x + 4, y + 4), s, fill='black', font=font)
            x += col_width
        y += row_h
    out = io.BytesIO()
    out_img.save(out, format='PNG')
    return out.getvalue()
