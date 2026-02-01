import sys
import io
import json
from time import sleep

try:
    import requests
except Exception:
    requests = None

try:
    from docx import Document
except Exception as e:
    print('python-docx missing:', e)
    sys.exit(1)

TEST_DOC = 'test_sensitive.docx'
REDACTED_OUT = 'redacted_test_sensitive.docx'

# create test docx
D = Document()
D.add_paragraph('Name: John Doe')
D.add_paragraph('SSN: 123-45-6789')
D.add_paragraph('Card: 4111 1111 1111 1111')
D.add_paragraph('Routing: 021000021')
D.add_paragraph('Sort code: 12-34-56')
D.add_paragraph('Password: secret123')
D.add_paragraph('OTP: 462345')
D.add_paragraph('Passport: A1234567')
D.add_paragraph('Diagnosis: patient has hypertension')
D.add_paragraph('Fingerprint: fingerprint sample')
D.save(TEST_DOC)
print('WROTE', TEST_DOC)

URL_BASE = 'http://127.0.0.1:8003'

# health check
try:
    if requests:
        h = requests.get(URL_BASE + '/health', timeout=5)
        print('HEALTH', h.status_code, h.text)
    else:
        import urllib.request
        r = urllib.request.urlopen(URL_BASE + '/health')
        print('HEALTH', r.status, r.read().decode())
except Exception as e:
    print('Health check failed:', e)
    sys.exit(1)

# /detect
print('\nPOST /detect')
try:
    if requests:
        r = requests.post(URL_BASE + '/detect', files={'file': open(TEST_DOC, 'rb')}, timeout=20)
        print('status', r.status_code)
        try:
            print(json.dumps(r.json(), indent=2))
        except Exception:
            print('non-json response', r.text[:200])
    else:
        import urllib.request, urllib.parse
        with open(TEST_DOC, 'rb') as f:
            data = f.read()
        boundary = '----WebKitFormBoundary7MA4YWxkTrZu0gW'
        body = []
        body.append('--' + boundary)
        body.append('Content-Disposition: form-data; name="file"; filename="%s"' % TEST_DOC)
        body.append('Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        body.append('')
        body.append(data.decode('ISO-8859-1'))
        body.append('--' + boundary + '--')
        req = urllib.request.Request(URL_BASE + '/detect', data='\r\n'.join(body).encode('ISO-8859-1'))
        req.add_header('Content-Type', 'multipart/form-data; boundary=%s' % boundary)
        resp = urllib.request.urlopen(req, timeout=20)
        print(resp.read().decode()[:1000])
except Exception as e:
    print('detect failed:', e)

# /redact/auto
print('\nPOST /redact/auto -> saving', REDACTED_OUT)
try:
    if requests:
        r = requests.post(URL_BASE + '/redact/auto', files={'file': open(TEST_DOC, 'rb')}, timeout=60)
        print('status', r.status_code, 'headers:', dict(r.headers))
        if r.status_code == 200:
            with open(REDACTED_OUT, 'wb') as out:
                out.write(r.content)
            print('Saved', REDACTED_OUT, 'size', str(os.path.getsize(REDACTED_OUT)) if 'os' in globals() else 'unknown')
        else:
            print('Redact returned:', r.text[:400])
    else:
        print('requests not installed; cannot POST easily')
except Exception as e:
    print('redact failed:', e)

print('\nDone')
